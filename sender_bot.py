"""
AEON Sender Bot - Redis Queue Sender for Bot API
Full support for documents generation from markdown (docx, pdf)
With Rate Limiting and Chat Actions for scalability (100+ → 5000+ users)

REACTIONS SUPPORT:
- Set reactions via Bot API set_message_reaction
- Remove reactions by passing empty emoji
- Works with PostgreSQL trigger notifications
"""

import threading
import time
import json
import asyncio
import logging
import redis
import psycopg2
import psycopg2.pool
import os
import re
import random
import hashlib
from typing import Optional, Dict, Any, Tuple
from io import BytesIO
from dataclasses import dataclass

from aiogram import Bot
from aiogram.types import (
    FSInputFile,
    BufferedInputFile,
    ReactionTypeEmoji,
    InputMediaPhoto,
    InputMediaVideo,
    InputMediaDocument
)
from aiogram.enums import ChatAction
from aiogram.exceptions import (
    TelegramRetryAfter,
    TelegramForbiddenError,
    TelegramBadRequest
)

# Database and Redis access layer
from db import db_cursor, db_connection, get_db_pool
from redis_client import get_redis, create_redis_client

# Message handlers
from handlers import MessageHandlers, HandlerResult

logger = logging.getLogger(__name__)


# ============== RATE LIMITER ==============

@dataclass
class RateLimitConfig:
    """Rate limit configuration for Bot API"""
    # Per-chat limits
    chat_messages_per_second: float = 1.0      # 1 msg/sec per chat
    group_messages_per_minute: int = 20        # 20 msg/min per group
    
    # Global limits  
    global_messages_per_second: int = 25       # 30 official, use 25 for safety
    global_api_requests_per_second: int = 25   # 30 official, use 25 for safety
    
    # Safety multiplier (reduce limits by this factor for extra safety)
    safety_factor: float = 0.85


class RedisRateLimiter:
    """
    Distributed rate limiter using Redis
    
    Implements sliding window rate limiting for:
    - Per-chat message limits
    - Per-group message limits  
    - Global message limits
    - Global API request limits
    
    Designed for horizontal scaling (multiple workers)
    """
    
    # Redis key prefixes
    PREFIX = "ratelimit:"
    CHAT_KEY = PREFIX + "chat:"           # Per-chat sliding window
    GROUP_KEY = PREFIX + "group:"         # Per-group sliding window
    GLOBAL_MSG_KEY = PREFIX + "global:msg"  # Global messages
    GLOBAL_API_KEY = PREFIX + "global:api"  # Global API requests
    
    def __init__(self, redis_client: redis.Redis, config: RateLimitConfig = None):
        self.redis = redis_client
        self.config = config or RateLimitConfig()
        
        # Stats
        self.stats = {
            'checks': 0,
            'passed': 0,
            'throttled_chat': 0,
            'throttled_group': 0,
            'throttled_global': 0,
            'wait_time_total': 0.0
        }
        
        logger.info(f"✅ RateLimiter initialized: {self.config}")
    
    def _is_group_chat(self, chat_id: int) -> bool:
        """Check if chat_id is a group/supergroup (negative ID)"""
        return chat_id < 0
    
    async def check_and_wait(self, chat_id: int, is_message: bool = True) -> float:
        """
        Check rate limits and wait if necessary
        
        Args:
            chat_id: Telegram chat ID
            is_message: True for messages, False for other API calls (typing, reactions, etc)
            
        Returns:
            Wait time in seconds (0 if no wait needed)
        """
        self.stats['checks'] += 1
        total_wait = 0.0
        now = time.time()
        
        try:
            # 1. Check global API limit (always)
            global_wait = await self._check_global_api(now)
            if global_wait > 0:
                self.stats['throttled_global'] += 1
                logger.debug(f"⏳ Global API throttle: {global_wait:.2f}s")
                await asyncio.sleep(global_wait)
                total_wait += global_wait
                now = time.time()
            
            # Record API request
            self._record_global_api(now)
            
            # For non-message API calls (typing, reactions, etc), only global limit applies
            if not is_message:
                self.stats['passed'] += 1
                return total_wait
            
            # 2. Check global message limit
            global_msg_wait = await self._check_global_messages(now)
            if global_msg_wait > 0:
                self.stats['throttled_global'] += 1
                logger.debug(f"⏳ Global msg throttle: {global_msg_wait:.2f}s")
                await asyncio.sleep(global_msg_wait)
                total_wait += global_msg_wait
                now = time.time()
            
            # 3. Check per-chat limit
            chat_wait = await self._check_chat_limit(chat_id, now)
            if chat_wait > 0:
                self.stats['throttled_chat'] += 1
                logger.debug(f"⏳ Chat {chat_id} throttle: {chat_wait:.2f}s")
                await asyncio.sleep(chat_wait)
                total_wait += chat_wait
                now = time.time()
            
            # 4. Check group limit (if applicable)
            if self._is_group_chat(chat_id):
                group_wait = await self._check_group_limit(chat_id, now)
                if group_wait > 0:
                    self.stats['throttled_group'] += 1
                    logger.debug(f"⏳ Group {chat_id} throttle: {group_wait:.2f}s")
                    await asyncio.sleep(group_wait)
                    total_wait += group_wait
                    now = time.time()
            
            # Record the message
            self._record_message(chat_id, now)
            
            self.stats['passed'] += 1
            self.stats['wait_time_total'] += total_wait
            
            return total_wait
            
        except redis.RedisError as e:
            logger.warning(f"⚠️ Rate limiter Redis error: {e}, allowing request")
            self.stats['passed'] += 1
            return 0.0
    
    async def _check_global_api(self, now: float) -> float:
        """Check global API requests limit"""
        key = self.GLOBAL_API_KEY
        window = 1.0  # 1 second window
        limit = int(self.config.global_api_requests_per_second * self.config.safety_factor)
        
        return self._sliding_window_check(key, now, window, limit)
    
    async def _check_global_messages(self, now: float) -> float:
        """Check global messages limit"""
        key = self.GLOBAL_MSG_KEY
        window = 1.0  # 1 second window
        limit = int(self.config.global_messages_per_second * self.config.safety_factor)
        
        return self._sliding_window_check(key, now, window, limit)
    
    async def _check_chat_limit(self, chat_id: int, now: float) -> float:
        """Check per-chat message limit"""
        key = f"{self.CHAT_KEY}{chat_id}"
        window = 1.0 / self.config.chat_messages_per_second  # Time between messages
        limit = 1
        
        return self._sliding_window_check(key, now, window, limit)
    
    async def _check_group_limit(self, chat_id: int, now: float) -> float:
        """Check per-group message limit (20/min)"""
        key = f"{self.GROUP_KEY}{chat_id}"
        window = 60.0  # 1 minute window
        limit = int(self.config.group_messages_per_minute * self.config.safety_factor)
        
        return self._sliding_window_check(key, now, window, limit)
    
    def _sliding_window_check(self, key: str, now: float, window: float, limit: int) -> float:
        """
        Sliding window rate limit check
        
        Returns wait time needed (0 if within limit)
        """
        # Get timestamps in window
        window_start = now - window
        
        # Use ZRANGEBYSCORE to get requests in window
        pipe = self.redis.pipeline()
        pipe.zremrangebyscore(key, 0, window_start)  # Clean old entries
        pipe.zcard(key)  # Count entries in window
        pipe.zrange(key, 0, 0, withscores=True)  # Get oldest entry
        results = pipe.execute()
        
        count = results[1]
        oldest = results[2]
        
        if count < limit:
            return 0.0
        
        # Calculate wait time
        if oldest:
            oldest_time = oldest[0][1]
            wait_time = (oldest_time + window) - now + 0.01  # Small buffer
            return max(0.0, wait_time)
        
        return window / limit  # Fallback wait
    
    def _record_global_api(self, now: float):
        """Record a global API request"""
        key = self.GLOBAL_API_KEY
        pipe = self.redis.pipeline()
        pipe.zadd(key, {f"{now}:{random.random()}": now})
        pipe.expire(key, 5)  # 5 second TTL
        pipe.execute()
    
    def _record_message(self, chat_id: int, now: float):
        """Record a message send"""
        pipe = self.redis.pipeline()
        
        # Global messages
        pipe.zadd(self.GLOBAL_MSG_KEY, {f"{now}:{chat_id}": now})
        pipe.expire(self.GLOBAL_MSG_KEY, 5)
        
        # Per-chat
        chat_key = f"{self.CHAT_KEY}{chat_id}"
        pipe.zadd(chat_key, {f"{now}": now})
        pipe.expire(chat_key, 5)
        
        # Per-group (if applicable)
        if self._is_group_chat(chat_id):
            group_key = f"{self.GROUP_KEY}{chat_id}"
            pipe.zadd(group_key, {f"{now}": now})
            pipe.expire(group_key, 120)  # 2 minute TTL
        
        pipe.execute()
    
    def get_stats(self) -> Dict:
        """Get rate limiter statistics"""
        return {
            **self.stats,
            'throttle_rate': (
                (self.stats['throttled_chat'] + self.stats['throttled_group'] + self.stats['throttled_global'])
                / max(1, self.stats['checks']) * 100
            ),
            'avg_wait': self.stats['wait_time_total'] / max(1, self.stats['passed'])
        }


# ============== CHAT ACTION MANAGER ==============

class ChatActionManager:
    """
    Manages chat actions (typing indicators) with deduplication
    
    Features:
    - Auto-selects correct action based on message type
    - Prevents spamming same action to same chat
    - Tracks action statistics
    """
    
    # Message type -> Chat action mapping
    ACTION_MAP = {
        'text': 'typing',
        'photo': 'upload_photo',
        'video': 'upload_video',
        'video_note': 'record_video_note',
        'voice': 'record_voice',
        'audio': 'upload_voice',
        'document': 'upload_document',
        'sticker': 'choose_sticker',
    }
    
    # ChatAction enum mapping
    CHAT_ACTION_ENUM = {
        'typing': ChatAction.TYPING,
        'upload_photo': ChatAction.UPLOAD_PHOTO,
        'upload_video': ChatAction.UPLOAD_VIDEO,
        'record_video_note': ChatAction.RECORD_VIDEO_NOTE,
        'record_video': ChatAction.RECORD_VIDEO,
        'record_voice': ChatAction.RECORD_VOICE,
        'upload_voice': ChatAction.UPLOAD_VOICE,
        'upload_document': ChatAction.UPLOAD_DOCUMENT,
        'choose_sticker': ChatAction.CHOOSE_STICKER,
    }
    
    def __init__(self, redis_client: redis.Redis, cooldown: float = 4.0):
        """
        Args:
            redis_client: Redis connection
            cooldown: Minimum seconds between same action to same chat
        """
        self.redis = redis_client
        self.cooldown = cooldown
        self.prefix = "chataction:"
        
        self.stats = {
            'sent': 0,
            'skipped_cooldown': 0,
            'skipped_no_action': 0,
            'errors': 0
        }
    
    def get_action_for_type(self, msg_type: str) -> Optional[str]:
        """Get chat action string for message type"""
        return self.ACTION_MAP.get(msg_type)
    
    def should_send_action(self, chat_id: int, action: str) -> bool:
        """
        Check if we should send this action (cooldown not expired)
        
        Returns True if action should be sent
        """
        key = f"{self.prefix}{chat_id}:{action}"
        
        # Try to set with NX (only if not exists)
        # Returns True if set (we should send), None if exists (skip)
        result = self.redis.set(key, "1", nx=True, ex=int(self.cooldown))
        
        return result is not None
    
    async def send_action(self, bot: Bot, chat_id: int, msg_type: str, 
                          rate_limiter: Optional[RedisRateLimiter] = None) -> bool:
        """
        Send appropriate chat action for message type
        
        Args:
            bot: Aiogram Bot instance
            chat_id: Telegram chat ID
            msg_type: Message type (text, photo, video, etc)
            rate_limiter: Optional rate limiter to check before sending
            
        Returns:
            True if action was sent, False otherwise
        """
        action = self.get_action_for_type(msg_type)
        
        if not action:
            self.stats['skipped_no_action'] += 1
            return False
        
        # Check cooldown
        if not self.should_send_action(chat_id, action):
            self.stats['skipped_cooldown'] += 1
            logger.debug(f"⏭️ Chat action cooldown: {chat_id} {action}")
            return False
        
        try:
            # Check rate limit if provided (action is API call but not message)
            if rate_limiter:
                await rate_limiter.check_and_wait(chat_id, is_message=False)
            
            chat_action = self.CHAT_ACTION_ENUM.get(action, ChatAction.TYPING)
            await bot.send_chat_action(chat_id=chat_id, action=chat_action)
            
            self.stats['sent'] += 1
            logger.debug(f"⌨️ Sent chat action: {chat_id} -> {action}")
            return True
            
        except TelegramForbiddenError:
            # User blocked bot or chat deleted - not an error
            self.stats['skipped_no_action'] += 1
            return False
            
        except Exception as e:
            self.stats['errors'] += 1
            logger.debug(f"⚠️ Chat action error ({chat_id}): {e}")
            return False
    
    def get_stats(self) -> Dict:
        """Get chat action statistics"""
        total = self.stats['sent'] + self.stats['skipped_cooldown'] + self.stats['skipped_no_action']
        return {
            **self.stats,
            'send_rate': self.stats['sent'] / max(1, total) * 100
        }


# ============== SENDER BOT ==============

class SenderBot:
    """
    Message sender for Bot API with document generation support
    
    Features:
    - Redis-based distributed rate limiting
    - Automatic chat actions (typing indicators)
    - Document generation (docx, pdf, txt)
    - All media types support
    - REACTIONS support (set/remove)
    - Horizontal scaling ready
    
    Supports:
    - text messages
    - stickers (by set name + emoji)
    - video notes (circles) with ffmpeg conversion
    - documents (generated from markdown: docx, pdf, txt)
    - reactions (set and remove)
    - photos, voice, video
    - delete messages
    """
    
    def __init__(self, bot: Bot, redis_url: str = None, db_url: str = None,
                 rate_limit_config: RateLimitConfig = None,
                 chat_action_cooldown: float = 4.0):
        self.bot = bot
        self.redis_url = redis_url  # Keep for reconnect compatibility
        self.redis_client = get_redis()

        self.queue_key = 'telegram:send_queue'
        self.sent_set_prefix = 'telegram:sent:'
        self.dedup_prefix = 'tg:dedup:'

        self.is_running = False
        self.worker_thread = None
        self.loop = None
        self.worker_bot = None

        # Rate Limiter
        self.rate_limiter = RedisRateLimiter(
            self.redis_client,
            rate_limit_config or RateLimitConfig()
        )

        # Chat Action Manager
        self.chat_action_manager = ChatActionManager(
            self.redis_client,
            cooldown=chat_action_cooldown
        )

        # Use shared database pool
        self.db_pool = get_db_pool()
        if self.db_pool:
            logger.info("Sender using shared DB pool")
        
        # In-memory set of successfully sent chat_history_ids
        # Survives DB failures — prevents duplicates even if _update_db_success fails
        # Only resets on process restart (acceptable: retry_stuck will re-check DB)
        self._sent_ids = set()
        self._sent_ids_lock = threading.Lock()

        # Stats
        self.stats = {
            'processed': 0,
            'sent': 0,
            'failed': 0,
            'retried': 0,
            'duplicates': 0,
            'documents_sent': 0,
            'video_notes_sent': 0,
            'reactions_sent': 0,
            'reactions_removed': 0,
            'stickers_sent': 0,
            'messages_deleted': 0
        }
        
        logger.info("✅ SenderBot initialized with rate limiting, chat actions, and reactions support")
    
    def start(self):
        """Start the worker thread"""
        if self.is_running:
            return
        
        self.is_running = True
        self.loop = asyncio.new_event_loop()
        self.worker_thread = threading.Thread(
            target=self._worker_loop,
            daemon=True,
            name="SenderBotWorker"
        )
        self.worker_thread.start()
        logger.info("✅ SenderBot worker started")
    
    def stop(self, timeout=10):
        """Stop the worker thread"""
        if not self.is_running:
            return

        logger.info("[SENDER] Stopping SenderBot...")
        self.is_running = False

        if self.worker_thread:
            self.worker_thread.join(timeout=timeout)

        # Don't close db_pool here — it's a shared pool managed by db.close_db_pool()

        try:
            stats = self.get_stats()
            logger.info(
                f"[SENDER] Stopped. Final stats: sent={stats.get('sent', 0)} "
                f"failed={stats.get('failed', 0)} retried={stats.get('retried', 0)}"
            )
        except Exception:
            logger.info("[SENDER] Stopped")
    
    def get_stats(self) -> Dict:
        """Get current statistics including rate limiter and chat actions"""
        return {
            **self.stats,
            'queue_length': self.redis_client.llen(self.queue_key),
            'rate_limiter': self.rate_limiter.get_stats(),
            'chat_actions': self.chat_action_manager.get_stats()
        }
    
    def queue_message(self, task_data: Dict[str, Any]):
        """Add message to Redis queue"""
        self.redis_client.rpush(self.queue_key, json.dumps(task_data))
        logger.debug(f"📤 Queued message: {task_data.get('chat_history_id')}")
    
    # ============== DOCUMENT GENERATION ==============
    
    def _generate_docx(self, markdown_text: str, filename: str = "document.docx") -> tuple:
        """
        Generate .docx from markdown with FULL formatting support
        
        Supports:
        - Headers (# ## ###)
        - Bold (**text**)
        - Italic (*text*)
        - Strikethrough (~~text~~)
        - Code (`code`)
        - Bullet lists (- item)
        - Numbered lists (1. item)
        """
        try:
            from docx import Document
            from docx.shared import RGBColor
            
            # Normalize newlines
            markdown_text = markdown_text.replace('\\n', '\n')
            
            doc = Document()
            lines = markdown_text.split('\n')
            
            def parse_inline_formatting(text, paragraph):
                """Parse inline markdown: **bold**, *italic*, `code`, ~~strike~~"""
                pattern = re.compile(
                    r'(\*\*\*(?P<bold_italic>[^\*]+)\*\*\*)|'
                    r'(\*\*(?P<bold>[^\*]+)\*\*)|'
                    r'(\*(?P<italic>[^*]+)\*)|'
                    r'(~~(?P<strike>[^~]+)~~)|'
                    r'(`(?P<code>[^`]+)`)|'
                    r'(?P<text>[^*~`]+)'
                )
                
                for match in pattern.finditer(text):
                    run = None
                    
                    if match.group('bold_italic'):
                        run = paragraph.add_run(match.group('bold_italic'))
                        run.bold = True
                        run.italic = True
                    elif match.group('bold'):
                        run = paragraph.add_run(match.group('bold'))
                        run.bold = True
                    elif match.group('italic'):
                        run = paragraph.add_run(match.group('italic'))
                        run.italic = True
                    elif match.group('strike'):
                        run = paragraph.add_run(match.group('strike'))
                        run.font.strike = True
                    elif match.group('code'):
                        run = paragraph.add_run(match.group('code'))
                        run.font.name = 'Courier New'
                        run.font.color.rgb = RGBColor(220, 50, 47)
                    elif match.group('text'):
                        run = paragraph.add_run(match.group('text'))
            
            for line in lines:
                line_stripped = line.strip()
                
                # Empty lines
                if not line_stripped:
                    doc.add_paragraph()
                    continue
                
                # Headers
                if line_stripped.startswith('# '):
                    para = doc.add_heading(level=1)
                    parse_inline_formatting(line_stripped[2:], para)
                elif line_stripped.startswith('## '):
                    para = doc.add_heading(level=2)
                    parse_inline_formatting(line_stripped[3:], para)
                elif line_stripped.startswith('### '):
                    para = doc.add_heading(level=3)
                    parse_inline_formatting(line_stripped[4:], para)
                
                # Bullet lists
                elif line_stripped.startswith('- ') or line_stripped.startswith('* '):
                    para = doc.add_paragraph(style='List Bullet')
                    parse_inline_formatting(line_stripped[2:], para)
                
                # Numbered lists
                elif re.match(r'^\d+\. ', line_stripped):
                    para = doc.add_paragraph(style='List Number')
                    text = line_stripped.split('. ', 1)[1]
                    parse_inline_formatting(text, para)
                
                # Regular text
                else:
                    para = doc.add_paragraph()
                    parse_inline_formatting(line_stripped, para)
            
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            logger.info(f"✅ DOCX generated: {len(buffer.getvalue())} bytes")
            return buffer, filename
            
        except Exception as e:
            logger.error(f"❌ DOCX generation failed: {e}")
            raise
    
    def _generate_pdf(self, markdown_text: str, filename: str = "document.pdf") -> tuple:
        """
        Generate .pdf from markdown with FULL formatting support
        
        Uses DejaVu fonts for Cyrillic support
        """
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            
            # Register Cyrillic fonts with proper family mapping
            font_name = 'Helvetica'
            font_name_bold = 'Helvetica-Bold'
            font_name_italic = 'Helvetica-Oblique'
            font_name_bolditalic = 'Helvetica-BoldOblique'
            
            try:
                # Check if fonts exist
                dejavu_regular = '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf'
                dejavu_bold = '/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf'
                dejavu_italic = '/usr/share/fonts/truetype/dejavu/DejaVuSans-Oblique.ttf'
                dejavu_bolditalic = '/usr/share/fonts/truetype/dejavu/DejaVuSans-BoldOblique.ttf'
                
                if os.path.exists(dejavu_regular):
                    pdfmetrics.registerFont(TTFont('DejaVuSans', dejavu_regular))
                    font_name = 'DejaVuSans'
                    logger.info("✅ DejaVuSans registered")
                    
                if os.path.exists(dejavu_bold):
                    pdfmetrics.registerFont(TTFont('DejaVuSans-Bold', dejavu_bold))
                    font_name_bold = 'DejaVuSans-Bold'
                    logger.info("✅ DejaVuSans-Bold registered")
                    
                if os.path.exists(dejavu_italic):
                    pdfmetrics.registerFont(TTFont('DejaVuSans-Oblique', dejavu_italic))
                    font_name_italic = 'DejaVuSans-Oblique'
                    logger.info("✅ DejaVuSans-Oblique registered")
                    
                if os.path.exists(dejavu_bolditalic):
                    pdfmetrics.registerFont(TTFont('DejaVuSans-BoldOblique', dejavu_bolditalic))
                    font_name_bolditalic = 'DejaVuSans-BoldOblique'
                    logger.info("✅ DejaVuSans-BoldOblique registered")
                
                # Register font family for automatic bold/italic mapping
                if font_name == 'DejaVuSans':
                    from reportlab.pdfbase.pdfmetrics import registerFontFamily
                    registerFontFamily(
                        'DejaVuSans',
                        normal='DejaVuSans',
                        bold='DejaVuSans-Bold',
                        italic='DejaVuSans-Oblique',
                        boldItalic='DejaVuSans-BoldOblique'
                    )
                    logger.info("✅ DejaVuSans font family registered")
                    
            except Exception as e:
                logger.warning(f"⚠️ Font registration error: {e}, using Helvetica")
                font_name = 'Helvetica'
                font_name_bold = 'Helvetica-Bold'
                font_name_italic = 'Helvetica-Oblique'
            
            buffer = BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=letter)
            styles = getSampleStyleSheet()
            
            # Create custom styles with Cyrillic font
            styles['Normal'].fontName = font_name
            styles['Normal'].fontSize = 11
            styles['Normal'].leading = 14
            
            styles['Heading1'].fontName = font_name_bold
            styles['Heading1'].fontSize = 18
            styles['Heading1'].leading = 22
            styles['Heading1'].textColor = '#1a365d'
            
            styles['Heading2'].fontName = font_name_bold
            styles['Heading2'].fontSize = 14
            styles['Heading2'].leading = 18
            styles['Heading2'].textColor = '#2c5282'
            
            styles['Heading3'].fontName = font_name_bold
            styles['Heading3'].fontSize = 12
            styles['Heading3'].leading = 16
            
            story = []
            
            # Normalize newlines
            markdown_text = markdown_text.replace('\\n', '\n')
            
            def parse_inline_to_html(text):
                """Convert markdown to HTML-like tags for reportlab"""
                text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                # Bold+Italic
                text = re.sub(r'\*\*\*([^\*]+)\*\*\*', rf'<font name="{font_name_bolditalic}">\1</font>', text)
                # Bold
                text = re.sub(r'\*\*([^\*]+)\*\*', rf'<font name="{font_name_bold}">\1</font>', text)
                # Italic
                text = re.sub(r'\*([^*]+)\*', rf'<font name="{font_name_italic}">\1</font>', text)
                # Strikethrough (reportlab doesn't support, just remove markers)
                text = re.sub(r'~~([^~]+)~~', r'\1', text)
                # Code - use monospace
                text = re.sub(r'`([^`]+)`', r'<font name="Courier" color="#c53030">\1</font>', text)
                return text
            
            lines = markdown_text.split('\n')
            
            for line in lines:
                line = line.strip()
                if not line:
                    story.append(Spacer(1, 0.15 * inch))
                    continue
                
                # Headers
                if line.startswith('# '):
                    text_html = parse_inline_to_html(line[2:])
                    story.append(Paragraph(text_html, styles['Heading1']))
                elif line.startswith('## '):
                    text_html = parse_inline_to_html(line[3:])
                    story.append(Paragraph(text_html, styles['Heading2']))
                elif line.startswith('### '):
                    text_html = parse_inline_to_html(line[4:])
                    story.append(Paragraph(text_html, styles['Heading3']))
                
                # Bullet lists
                elif line.startswith('- ') or line.startswith('* '):
                    text_html = parse_inline_to_html(line[2:])
                    story.append(Paragraph(f"• {text_html}", styles['Normal']))
                
                # Numbered lists
                elif re.match(r'^\d+\. ', line):
                    match = re.match(r'^(\d+)\. (.+)', line)
                    if match:
                        num, text = match.groups()
                        text_html = parse_inline_to_html(text)
                        story.append(Paragraph(f"{num}. {text_html}", styles['Normal']))
                
                # Regular text
                else:
                    text_html = parse_inline_to_html(line)
                    story.append(Paragraph(text_html, styles['Normal']))
            
            doc.build(story)
            buffer.seek(0)
            
            logger.info(f"✅ PDF generated: {len(buffer.getvalue())} bytes")
            return buffer, filename
            
        except Exception as e:
            logger.error(f"❌ PDF generation failed: {e}")
            raise
    
    def _generate_txt(self, markdown_text: str, filename: str = "document.txt") -> tuple:
        """Generate .txt from markdown (just clean text)"""
        try:
            # Normalize newlines
            text = markdown_text.replace('\\n', '\n')
            
            # Remove markdown formatting for plain text
            # Remove headers markers
            text = re.sub(r'^#{1,6}\s+', '', text, flags=re.MULTILINE)
            # Remove bold/italic markers
            text = re.sub(r'\*\*\*([^\*]+)\*\*\*', r'\1', text)
            text = re.sub(r'\*\*([^\*]+)\*\*', r'\1', text)
            text = re.sub(r'\*([^*]+)\*', r'\1', text)
            # Remove strikethrough
            text = re.sub(r'~~([^~]+)~~', r'\1', text)
            # Remove code markers
            text = re.sub(r'`([^`]+)`', r'\1', text)
            
            buffer = BytesIO()
            buffer.write(text.encode('utf-8'))
            buffer.seek(0)
            
            logger.info(f"✅ TXT generated: {len(buffer.getvalue())} bytes")
            return buffer, filename
            
        except Exception as e:
            logger.error(f"❌ TXT generation failed: {e}")
            raise
    
    # ============== SEND METHODS ==============
    
    def _markdown_to_html(self, text: str) -> str:
        """
        Convert Markdown to HTML for Telegram

        Supports:
        - **bold** -> <b>bold</b>
        - *italic* -> <i>italic</i>
        - `code` -> <code>code</code>
        - [text](url) -> <a href="url">text</a>
        - ~~strike~~ -> <s>strike</s>
        """
        import html

        # Escape HTML special chars first (but preserve our markdown)
        # We need to be careful not to break markdown syntax

        # Process links first (before escaping) - [text](url)
        def replace_link(match):
            link_text = match.group(1)
            url = match.group(2)
            # Escape the link text, but not the URL
            return f'<a href="{url}">{html.escape(link_text)}</a>'

        text = re.sub(r'\[([^\]]+)\]\(([^)]+)\)', replace_link, text)

        # Mark placeholders for markdown we want to preserve
        # Bold: **text**
        bold_pattern = re.compile(r'\*\*([^\*]+)\*\*')
        italic_pattern = re.compile(r'(?<!\*)\*([^*]+)\*(?!\*)')
        code_pattern = re.compile(r'`([^`]+)`')
        strike_pattern = re.compile(r'~~([^~]+)~~')

        # Extract and replace with placeholders
        placeholders = {}
        counter = [0]

        def make_placeholder(match, tag_open, tag_close):
            key = f"__PH{counter[0]}__"
            counter[0] += 1
            content = html.escape(match.group(1))
            placeholders[key] = f"{tag_open}{content}{tag_close}"
            return key

        text = bold_pattern.sub(lambda m: make_placeholder(m, '<b>', '</b>'), text)
        text = italic_pattern.sub(lambda m: make_placeholder(m, '<i>', '</i>'), text)
        text = code_pattern.sub(lambda m: make_placeholder(m, '<code>', '</code>'), text)
        text = strike_pattern.sub(lambda m: make_placeholder(m, '<s>', '</s>'), text)

        # Now escape remaining HTML chars (links already have escaped text)
        # But don't escape our <a> tags
        parts = re.split(r'(<a href="[^"]+">.*?</a>)', text)
        escaped_parts = []
        for part in parts:
            if part.startswith('<a href='):
                escaped_parts.append(part)
            else:
                escaped_parts.append(html.escape(part))
        text = ''.join(escaped_parts)

        # Restore placeholders
        for key, value in placeholders.items():
            text = text.replace(html.escape(key), value)

        return text

    async def _send_text(self, chat_id: int, text: str, reply_to: int = None,
                         message_thread_id: int = None) -> int:
        """Send text message with HTML (converted from Markdown), fallback to plain text"""
        from aiogram.enums import ParseMode

        # Normalize newlines
        text = text.replace('\\n', '\n')

        # Convert Markdown to HTML for better link/formatting support
        html_text = self._markdown_to_html(text)

        # Try with HTML first (more reliable for links)
        try:
            msg = await self.worker_bot.send_message(
                chat_id=chat_id,
                text=html_text,
                reply_to_message_id=reply_to,
                message_thread_id=message_thread_id,
                parse_mode=ParseMode.HTML
            )
            return msg.message_id
        except TelegramBadRequest as e:
            error_msg = str(e).lower()
            if "can't parse" in error_msg or "parse" in error_msg:
                # Fallback to plain text
                logger.warning(f"⚠️ HTML parse failed, sending as plain text: {e}")
                msg = await self.worker_bot.send_message(
                    chat_id=chat_id,
                    text=text,  # Original text without HTML
                    reply_to_message_id=reply_to,
                    message_thread_id=message_thread_id,
                    parse_mode=None
                )
                return msg.message_id
            raise
    
    async def _send_sticker(self, chat_id: int, sticker_set_name: str = None,
                            emoji: str = None, sticker_file_id: str = None,
                            reply_to: int = None, message_thread_id: int = None) -> int:
        """
        Send sticker by set name + emoji or by file_id

        Args:
            chat_id: Telegram chat ID
            sticker_set_name: Name of sticker set (e.g. "AEON31")
            emoji: Emoji to find in sticker set (e.g. "👍")
            sticker_file_id: Direct file_id of sticker (priority over set+emoji)
            reply_to: Message ID to reply to
            message_thread_id: Forum topic thread ID (optional)
        """
        # Priority 1: Direct file_id
        if sticker_file_id:
            logger.info(f"📎 Sending sticker by file_id: {sticker_file_id[:20]}...")
            msg = await self.worker_bot.send_sticker(
                chat_id=chat_id,
                sticker=sticker_file_id,
                reply_to_message_id=reply_to,
                message_thread_id=message_thread_id
            )
            return msg.message_id
        
        # Priority 2: Set name + emoji
        if not sticker_set_name:
            raise ValueError("sticker_set_name or sticker_file_id required")
        
        logger.info(f"🎨 Getting sticker set: {sticker_set_name}, looking for emoji: {emoji}")
        
        # Get sticker set from Telegram
        sticker_set = await self.worker_bot.get_sticker_set(sticker_set_name)
        
        if not sticker_set.stickers:
            raise ValueError(f"Empty sticker set: {sticker_set_name}")
        
        logger.info(f"📦 Sticker set '{sticker_set_name}' has {len(sticker_set.stickers)} stickers")
        
        # Find sticker by emoji
        sticker = None
        if emoji:
            # Normalize emoji (remove variation selectors)
            emoji_normalized = emoji.strip()
            
            # Try multiple matching strategies
            for s in sticker_set.stickers:
                sticker_emoji = s.emoji or ""
                
                # Strategy 1: Exact match
                if sticker_emoji == emoji_normalized:
                    sticker = s
                    logger.info(f"✅ Found sticker by exact match: {sticker_emoji}")
                    break
                
                # Strategy 2: Emoji contains target (sticker may have multiple emoji)
                if emoji_normalized in sticker_emoji:
                    sticker = s
                    logger.info(f"✅ Found sticker by contains: {sticker_emoji} contains {emoji_normalized}")
                    break
                
                # Strategy 3: Target contains sticker emoji
                if sticker_emoji in emoji_normalized:
                    sticker = s
                    logger.info(f"✅ Found sticker by reverse contains: {emoji_normalized} contains {sticker_emoji}")
                    break
            
            if not sticker:
                # Log all available emojis for debugging
                available_emojis = [s.emoji for s in sticker_set.stickers[:10]]
                logger.warning(f"⚠️ Emoji '{emoji}' not found in set. Available (first 10): {available_emojis}")
        
        # Fallback to first sticker
        if not sticker:
            sticker = sticker_set.stickers[0]
            logger.info(f"📎 Using first sticker as fallback, emoji: {sticker.emoji}")
        
        # Send sticker
        msg = await self.worker_bot.send_sticker(
            chat_id=chat_id,
            sticker=sticker.file_id,
            reply_to_message_id=reply_to,
            message_thread_id=message_thread_id
        )
        
        logger.info(f"✅ Sticker sent: set={sticker_set_name}, emoji={sticker.emoji}, msg_id={msg.message_id}")
        return msg.message_id
    
    async def _send_video_note(self, chat_id: int, video_data: bytes,
                                reply_to: int = None, message_thread_id: int = None) -> int:
        """
        Send video note (circle video)
        
        Note: Bot API accepts video notes as-is, but they should be:
        - Square format (1:1 aspect ratio)
        - Max 1 minute duration
        - Max 640x640 pixels
        
        For non-square videos, consider using ffmpeg to crop before sending.
        """
        import tempfile
        import subprocess
        
        # Check if ffmpeg is available for cropping
        try:
            # Create temp files
            with tempfile.NamedTemporaryFile(suffix='.mp4', delete=False) as tmp_in:
                tmp_in.write(video_data)
                tmp_in_path = tmp_in.name
            
            tmp_out_path = tmp_in_path.replace('.mp4', '_square.mp4')
            
            try:
                # Crop to square with ffmpeg
                result = subprocess.run([
                    'ffmpeg', '-y', '-i', tmp_in_path,
                    '-vf', 'crop=min(iw\\,ih):min(iw\\,ih),scale=640:640',
                    '-c:v', 'libx264', '-crf', '28',
                    '-c:a', 'copy',
                    tmp_out_path
                ], capture_output=True, timeout=60)
                
                if result.returncode == 0:
                    with open(tmp_out_path, 'rb') as f:
                        video_data = f.read()
                    logger.info(f"✅ Video cropped to square: {len(video_data)} bytes")
                else:
                    logger.warning(f"⚠️ ffmpeg crop failed, using original video")
                    
            finally:
                # Cleanup
                try:
                    os.unlink(tmp_in_path)
                except:
                    pass
                try:
                    os.unlink(tmp_out_path)
                except:
                    pass
                    
        except FileNotFoundError:
            logger.warning("⚠️ ffmpeg not found, sending video as-is")
        except Exception as e:
            logger.warning(f"⚠️ Video processing error: {e}, sending as-is")
        
        # Send video note
        video_note = BufferedInputFile(video_data, filename="video_note.mp4")
        
        msg = await self.worker_bot.send_video_note(
            chat_id=chat_id,
            video_note=video_note,
            reply_to_message_id=reply_to,
            message_thread_id=message_thread_id
        )
        return msg.message_id

    async def _send_document_generated(self, chat_id: int, markdown_text: str,
                                       file_format: str = 'docx', filename: str = None,
                                       reply_to: int = None,
                                       message_thread_id: int = None) -> int:
        """
        Generate and send document from markdown text
        
        Args:
            chat_id: Telegram chat ID
            markdown_text: Markdown content for document
            file_format: 'docx', 'pdf', or 'txt'
            filename: Output filename (optional)
            reply_to: Reply to message ID
        """
        # Determine format and default filename
        file_format = file_format.lower().strip('.')
        
        if file_format == 'pdf':
            default_filename = filename or "document.pdf"
            file_buffer, fname = self._generate_pdf(markdown_text, default_filename)
        elif file_format == 'txt':
            default_filename = filename or "document.txt"
            file_buffer, fname = self._generate_txt(markdown_text, default_filename)
        else:  # docx by default
            default_filename = filename or "document.docx"
            file_buffer, fname = self._generate_docx(markdown_text, default_filename)
        
        # Send document
        document = BufferedInputFile(file_buffer.read(), filename=fname)
        
        msg = await self.worker_bot.send_document(
            chat_id=chat_id,
            document=document,
            reply_to_message_id=reply_to,
            message_thread_id=message_thread_id
        )

        logger.info(f"✅ Document sent: {fname}")
        return msg.message_id

    async def _send_document_file(self, chat_id: int, file_data: bytes,
                                  filename: str = None, caption: str = None,
                                  reply_to: int = None,
                                  message_thread_id: int = None) -> int:
        """Send ready document file (not generated)"""
        if not filename:
            filename = 'document'
        document = BufferedInputFile(file_data, filename=filename)

        msg = await self.worker_bot.send_document(
            chat_id=chat_id,
            document=document,
            caption=caption,
            reply_to_message_id=reply_to,
            message_thread_id=message_thread_id
        )
        return msg.message_id

    async def _send_reaction(self, chat_id: int, message_id: int, emoji: str) -> bool:
        """
        Set or remove reaction on message

        Args:
            chat_id: Telegram chat ID
            message_id: Message ID to react to
            emoji: Reaction emoji (None or empty to remove reaction)

        Returns:
            True if successful
        """
        # DEBUG: Log input
        logger.info(f"🔧 [REACTION] _send_reaction: chat={chat_id}, msg={message_id}, emoji={repr(emoji)}")

        # Build reaction list
        if emoji and emoji.strip():
            reaction = [ReactionTypeEmoji(emoji=emoji)]
            action = "set"
        else:
            reaction = []
            action = "remove"

        logger.info(f"🔧 [REACTION] Calling set_message_reaction with reaction={reaction}")

        await self.worker_bot.set_message_reaction(
            chat_id=chat_id,
            message_id=message_id,
            reaction=reaction
        )

        logger.info(f"✅ Reaction {action}: chat={chat_id}, msg={message_id}, emoji={emoji}")
        return True
    
    async def _delete_message(self, chat_id: int, message_id: int) -> bool:
        """
        Delete message from Telegram

        Args:
            chat_id: Telegram chat ID
            message_id: Message ID to delete (tg_id from chat_history_tg)

        Returns:
            True if deleted successfully
        """
        try:
            await self.worker_bot.delete_message(chat_id=chat_id, message_id=message_id)
            logger.info(f"🗑️ Message deleted: chat={chat_id}, msg_id={message_id}")
            return True
        except TelegramBadRequest as e:
            # Message already deleted or not found
            if "message to delete not found" in str(e).lower():
                logger.warning(f"⚠️ Message already deleted: chat={chat_id}, msg_id={message_id}")
                return True  # Consider it success - message is gone
            elif "message can't be deleted" in str(e).lower():
                logger.warning(f"⚠️ Message can't be deleted (too old?): chat={chat_id}, msg_id={message_id}")
                return False
            raise
        except TelegramForbiddenError:
            logger.warning(f"⚠️ No permission to delete message: chat={chat_id}, msg_id={message_id}")
            return False

    async def _send_typing(self, chat_id: int, action: str = 'typing'):
        """Send chat action (typing indicator)"""
        action_map = {
            'typing': ChatAction.TYPING,
            'upload_photo': ChatAction.UPLOAD_PHOTO,
            'upload_video': ChatAction.UPLOAD_VIDEO,
            'upload_document': ChatAction.UPLOAD_DOCUMENT,
            'record_video': ChatAction.RECORD_VIDEO,
            'record_video_note': ChatAction.RECORD_VIDEO_NOTE,
            'record_voice': ChatAction.RECORD_VOICE,
            'upload_voice': ChatAction.UPLOAD_VOICE,
            'choose_sticker': ChatAction.CHOOSE_STICKER,
        }
        
        chat_action = action_map.get(action, ChatAction.TYPING)
        await self.worker_bot.send_chat_action(chat_id=chat_id, action=chat_action)
    
    async def _send_photo(self, chat_id: int, photo_data: bytes,
                          caption: str = None, reply_to: int = None,
                          message_thread_id: int = None) -> int:
        """Send photo"""
        photo = BufferedInputFile(photo_data, filename="photo.jpg")

        msg = await self.worker_bot.send_photo(
            chat_id=chat_id,
            photo=photo,
            caption=caption,
            reply_to_message_id=reply_to,
            message_thread_id=message_thread_id
        )
        return msg.message_id
    
    async def _send_voice(self, chat_id: int, voice_data: bytes,
                          reply_to: int = None, message_thread_id: int = None) -> int:
        """Send voice message"""
        voice = BufferedInputFile(voice_data, filename="voice.ogg")

        msg = await self.worker_bot.send_voice(
            chat_id=chat_id,
            voice=voice,
            reply_to_message_id=reply_to,
            message_thread_id=message_thread_id
        )
        return msg.message_id
    
    async def _send_video(self, chat_id: int, video_data: bytes,
                          caption: str = None, reply_to: int = None,
                          message_thread_id: int = None) -> int:
        """Send video"""
        video = BufferedInputFile(video_data, filename="video.mp4")

        msg = await self.worker_bot.send_video(
            chat_id=chat_id,
            video=video,
            caption=caption,
            reply_to_message_id=reply_to,
            message_thread_id=message_thread_id
        )
        return msg.message_id

    async def _send_media_group(self, chat_id: int, media: list,
                                caption: str = None, reply_to: int = None,
                                message_thread_id: int = None) -> list:
        """
        Send media group (album) - up to 10 photos/videos

        Args:
            chat_id: Telegram chat ID
            media: list of dicts [{type: "photo"|"video", url: "..."}, ...]
            caption: Caption for first media item
            reply_to: Message ID to reply to

        Returns:
            list of message IDs for each sent media
        """
        if not media:
            return []

        media_list = []
        for i, item in enumerate(media):
            media_type = item.get('type', 'photo')
            url = item.get('url')

            if not url:
                continue

            # Caption only on first item
            item_caption = caption if i == 0 else None

            if media_type == 'photo':
                media_list.append(InputMediaPhoto(
                    media=url,
                    caption=item_caption
                ))
            elif media_type == 'video':
                media_list.append(InputMediaVideo(
                    media=url,
                    caption=item_caption
                ))
            elif media_type == 'document':
                media_list.append(InputMediaDocument(
                    media=url,
                    caption=item_caption
                ))

        if len(media_list) < 2:
            logger.warning(f"Media group needs at least 2 items, got {len(media_list)}")
            return []

        messages = await self.worker_bot.send_media_group(
            chat_id=chat_id,
            media=media_list,
            reply_to_message_id=reply_to,
            message_thread_id=message_thread_id
        )

        return [msg.message_id for msg in messages]

    # ============== TASK PROCESSING ==============
    
    async def _async_process_task(self, task_data: Dict[str, Any]):
        """Process single task from queue"""
        chat_history_id = task_data.get('chat_history_id')
        request_body = task_data.get('request_body', {})
        retry_count = task_data.get('retry_count', 0)
        
        # Get telegram_chat_id from multiple sources:
        # 1. task_data.telegram_chat_id (from API endpoint or new trigger)
        # 2. task_data.chat_ident (from PostgreSQL trigger - legacy, should be telegram_id)
        # 3. request_body.chat_id (fallback)
        telegram_chat_id = task_data.get('telegram_chat_id')
        if not telegram_chat_id:
            telegram_chat_id = task_data.get('chat_ident')
        if not telegram_chat_id:
            telegram_chat_id = request_body.get('chat_id')
        
        # Convert to int (Bot API requires numeric chat_id)
        if isinstance(telegram_chat_id, str):
            # Check if it's a valid numeric string
            try:
                telegram_chat_id = int(telegram_chat_id)
            except ValueError:
                # Not a numeric string - this could be username, which Bot API doesn't support
                logger.error(f"❌ Invalid telegram_chat_id (not numeric): {telegram_chat_id}")
                logger.error(f"   Bot API requires numeric telegram_id, not username/phone")
                self._handle_failure(chat_history_id, "Invalid telegram_chat_id: not numeric")
                return

        if not telegram_chat_id:
            logger.error(f"❌ No telegram_chat_id in task: {task_data}")
            self._handle_failure(chat_history_id, "No telegram_chat_id in task")
            return
        
        self.stats['processed'] += 1
        
        msg_type = request_body.get('type_of_message', 'text')
        is_reaction = msg_type == 'reaction'
        is_typing_only = msg_type == 'typing'
        is_delete = msg_type == 'delete'
        
        # ========================================
        # 🔥 ДЕДУПЛИКАЦИЯ (как в PostgreSQL индексе)
        # chat_id + type + type_of_message + thought_sessions_id + md5(message) + 5-min slot
        # ========================================
        skip_dedup = task_data.get('force_send', False)
        if not is_reaction and not is_typing_only and not is_delete and not skip_dedup:
            
            thought_sessions_id = task_data.get('thought_sessions_id')
            message_text = request_body.get('message', '')
            
            # Дедупликация по контенту (5-минутное окно)
            if thought_sessions_id and message_text:
                msg_hash = hashlib.md5(message_text.encode()).hexdigest()
                time_slot = int(time.time()) // 300  # 5-минутный слот (300 сек)
                
                # Ключ: chat:type:type_of_message:session:hash:slot
                dedup_key = f"{self.dedup_prefix}{telegram_chat_id}:AEON:{msg_type}:{thought_sessions_id}:{msg_hash}:{time_slot}"
                
                if not self.redis_client.set(dedup_key, str(chat_history_id or ''), nx=True, ex=600):
                    self.stats['duplicates'] += 1
                    logger.warning(
                        f"⚠️ Duplicate [content]: chat={telegram_chat_id}, "
                        f"session={thought_sessions_id}, type={msg_type}, hash={msg_hash[:8]}"
                    )
                    return
            
            # Базовая дедупликация по chat_history_id (fallback)
            if chat_history_id:
                lock_key = f"{self.sent_set_prefix}{chat_history_id}"
                if not self.redis_client.set(lock_key, '1', nx=True, ex=600):
                    self.stats['duplicates'] += 1
                    logger.warning(f"⚠️ Duplicate [id]: {chat_history_id}")
                    return

        # ========== PRE-SEND DEDUP (3 layers) ==========
        if chat_history_id and not is_reaction and not is_typing_only and not is_delete:
            # Layer 1: In-memory set (survives DB failures, fastest check)
            with self._sent_ids_lock:
                if chat_history_id in self._sent_ids:
                    logger.info(f"⏭️ Already sent (in-memory), skipping: chat_history_id={chat_history_id}")
                    self.stats['duplicates'] += 1
                    return

            # Layer 2: DB check (catches retries after process restart)
            try:
                with db_cursor() as cur:
                    cur.execute("SELECT tg_id, status FROM chat_history_tg WHERE id = %s", (chat_history_id,))
                    row = cur.fetchone()
                    if row and row[0] is not None:
                        logger.info(f"⏭️ Already sent (tg_id={row[0]}), skipping: chat_history_id={chat_history_id}")
                        self.stats['duplicates'] += 1
                        with self._sent_ids_lock:
                            self._sent_ids.add(chat_history_id)
                        return
                    if row and row[1] == 'sent':
                        logger.info(f"⏭️ Status=sent but no tg_id, skipping: chat_history_id={chat_history_id}")
                        self.stats['duplicates'] += 1
                        return
            except Exception as e:
                logger.debug(f"Pre-send DB check failed (proceeding): {e}")

        try:
            # Stage timing — surfaces where latency is spent (rate-limit, typing, send, DB).
            _t0 = time.time()
            queued_at = task_data.get('queued_at')
            queue_wait = (_t0 - float(queued_at)) if queued_at else None

            # ========== RATE LIMITING ==========
            # Check rate limits before any action (except for typing-only and delete requests)
            if not is_typing_only and not is_delete:
                wait_time = await self.rate_limiter.check_and_wait(
                    telegram_chat_id,
                    is_message=(not is_reaction)  # Reactions are API calls, not messages
                )
                if wait_time > 0:
                    logger.debug(f"⏳ Rate limited {telegram_chat_id}: waited {wait_time:.2f}s")
            _t1 = time.time()

            # ========== CHAT ACTION ==========
            # Send typing/uploading indicator before actual send
            # Skip for reactions, typing-only, and delete requests
            show_typing = request_body.get('show_typing', True)
            if show_typing and not is_reaction and not is_typing_only and not is_delete:
                await self.chat_action_manager.send_action(
                    self.worker_bot,
                    telegram_chat_id,
                    msg_type,
                    self.rate_limiter
                )
            _t2 = time.time()

            # ========== MESSAGE SENDING via HANDLERS ==========
            result: HandlerResult = await MessageHandlers.handle(
                self, telegram_chat_id, msg_type, request_body
            )
            _t3 = time.time()

            if result.success:
                self.stats['sent'] += 1
                if result.stat_key:
                    self.stats[result.stat_key] += 1

                stage_summary = (
                    f"id={chat_history_id} type={msg_type} tg_id={result.tg_message_id} "
                    f"queue_wait={queue_wait:.2f}s " if queue_wait is not None else f"id={chat_history_id} type={msg_type} tg_id={result.tg_message_id} "
                )
                stage_summary += (
                    f"ratelimit={_t1-_t0:.2f}s typing={_t2-_t1:.2f}s "
                    f"send={_t3-_t2:.2f}s"
                )
                if (_t3 - _t0) > 3.0 or (queue_wait is not None and queue_wait > 3.0):
                    logger.warning(f"[SLOW SEND] {stage_summary}")
                else:
                    logger.info(f"Sent {stage_summary}")

                # Mark as sent in memory IMMEDIATELY (before DB update)
                # This prevents duplicates even if DB update fails
                if chat_history_id:
                    with self._sent_ids_lock:
                        self._sent_ids.add(chat_history_id)
                        # Cap set size to prevent memory leak
                        if len(self._sent_ids) > 50000:
                            # Remove oldest ~10k entries (set is unordered, but that's fine)
                            to_remove = list(self._sent_ids)[:10000]
                            self._sent_ids -= set(to_remove)

                # Update DB status
                if chat_history_id and result.tg_message_id:
                    self._update_db_success(chat_history_id, result.tg_message_id)
            else:
                if result.error:
                    raise ValueError(result.error)

        except TelegramRetryAfter as e:
            logger.warning(f"⏳ Rate limit from Telegram: wait {e.retry_after}s")
            
            # Update rate limiter awareness
            await asyncio.sleep(e.retry_after + 1)
            
            if retry_count < 3:
                task_data['retry_count'] = retry_count + 1
                self.redis_client.rpush(self.queue_key, json.dumps(task_data))
                self.stats['retried'] += 1
            else:
                self._handle_failure(chat_history_id, str(e))
        
        except TelegramForbiddenError as e:
            logger.error(f"❌ Forbidden: {e}")
            self._handle_failure(chat_history_id, f"Forbidden: {e}")
        
        except TelegramBadRequest as e:
            logger.error(f"❌ Bad request: {e}")
            self._handle_failure(chat_history_id, f"BadRequest: {e}")
        
        except Exception as e:
            logger.exception(f"[SENDER] Send error chat_history_id={chat_history_id}: {e}")

            # Only retry if NOT already sent (error could be from DB update after successful TG send)
            already_sent = False
            if chat_history_id:
                with self._sent_ids_lock:
                    already_sent = chat_history_id in self._sent_ids

            if already_sent:
                logger.info(f"⏭️ Error after successful send, NOT retrying: chat_history_id={chat_history_id}")
            elif retry_count < 3:
                task_data['retry_count'] = retry_count + 1
                # DO NOT set force_send — retry must go through all dedup checks
                task_data['retry_after'] = time.time() + (2 ** retry_count) * 5
                self.redis_client.rpush(self.queue_key, json.dumps(task_data))
                self.stats['retried'] += 1
            else:
                self._handle_failure(chat_history_id, str(e))
    
    def _handle_failure(self, chat_history_id: int, error: str):
        """Handle failed message"""
        self.stats['failed'] += 1
        
        if chat_history_id:
            self._update_db_failed(chat_history_id, error)
            self.redis_client.delete(f"{self.sent_set_prefix}{chat_history_id}")
    
    def _update_db_success(self, chat_history_id: int, tg_message_id: int):
        """Update DB on successful send"""
        if not get_db_pool():
            return

        for attempt in range(3):
            try:
                with db_cursor(commit=True) as cur:
                    try:
                        cur.execute(
                            "SELECT update_message_status(%s, %s, %s, %s)",
                            (chat_history_id, tg_message_id, 'sent', None)
                        )
                    except Exception:
                        cur.execute("""
                            UPDATE chat_history_tg
                            SET tg_id = %s, status = 'sent'
                            WHERE id = %s
                        """, (tg_message_id, chat_history_id))
                return  # success
            except psycopg2.pool.PoolError:
                if attempt < 2:
                    time.sleep(0.5 * (attempt + 1))
            except Exception as e:
                logger.error(f"DB update error: {e}")
                return

    def _update_db_failed(self, chat_history_id: int, error: str):
        """Update DB on failed send"""
        if not get_db_pool():
            return

        for attempt in range(3):
            try:
                with db_cursor(commit=True) as cur:
                    try:
                        cur.execute(
                            "SELECT update_message_status(%s, %s, %s, %s)",
                            (chat_history_id, None, 'failed', error[:500])
                        )
                    except Exception:
                        cur.execute("""
                            UPDATE chat_history_tg
                            SET status = 'failed', tg_error_message = %s
                            WHERE id = %s
                        """, (error, chat_history_id,))
                return  # success
            except psycopg2.pool.PoolError:
                if attempt < 2:
                    time.sleep(0.5 * (attempt + 1))
            except Exception as e:
                logger.error(f"DB failed update error: {e}")
                return
    
    def _worker_loop(self):
        """Main worker loop - runs async event loop"""
        logger.info("🔄 SenderBot worker loop started")
        
        self.loop = asyncio.new_event_loop()
        asyncio.set_event_loop(self.loop)
        
        from aiogram.client.default import DefaultBotProperties
        from aiogram.enums import ParseMode
        
        bot_token = self.bot.token
        # No default parse_mode - we handle it per-message with fallback
        self.worker_bot = Bot(token=bot_token)
        logger.info("✅ Worker Bot instance created")
        
        try:
            # Create task explicitly - required for aiohttp 3.9+ timeout handling
            task = self.loop.create_task(self._async_worker())
            self.loop.run_until_complete(task)
        except Exception as e:
            logger.error(f"SenderBot worker loop crashed: {e}")
        finally:
            try:
                self.loop.run_until_complete(self.worker_bot.session.close())
            except:
                pass
            self.loop.close()
            self.is_running = False  # Signal that worker is dead

        logger.info("✅ SenderBot worker loop stopped")
    
    async def _async_worker(self):
        """Async worker that processes Redis queue"""
        consecutive_errors = 0
        
        while self.is_running:
            try:
                result = self.redis_client.blpop(self.queue_key, timeout=2)
                
                if not result:
                    continue
                
                _, task_json = result
                task_data = json.loads(task_json)

                # DEBUG: Log incoming task for reactions
                request_body = task_data.get('request_body', {})
                if request_body.get('type_of_message') == 'reaction':
                    logger.info(f"📥 [REACTION DEBUG] Received task: {task_json}")

                # Check retry delay
                retry_after = task_data.get('retry_after', 0)
                if retry_after > time.time():
                    wait = retry_after - time.time()
                    logger.debug(f"⏰ Delayed task: {wait:.1f}s")
                    await asyncio.sleep(wait)
                
                consecutive_errors = 0
                await self._async_process_task(task_data)
                
            except json.JSONDecodeError as e:
                logger.error(f"❌ JSON error: {e}")
                consecutive_errors += 1
                
            except redis.ConnectionError as e:
                consecutive_errors += 1
                if consecutive_errors == 1 or consecutive_errors % 10 == 0:
                    logger.error(f"[SENDER] Redis connection error (#{consecutive_errors}): {e}")
                backoff = min(2 ** min(consecutive_errors, 5), 30)
                await asyncio.sleep(backoff)

                if consecutive_errors > 5:
                    try:
                        old_client = self.redis_client
                        self.redis_client = create_redis_client()
                        if self.redis_client.ping():
                            logger.info("[SENDER] Redis reconnected")
                            consecutive_errors = 0
                        # Close old client if it's a different instance
                        if old_client and old_client is not self.redis_client:
                            try:
                                old_client.close()
                            except Exception:
                                pass
                    except Exception:
                        await asyncio.sleep(5)
                        
            except Exception as e:
                consecutive_errors += 1
                if consecutive_errors == 1 or consecutive_errors % 5 == 0:
                    logger.exception(f"[SENDER] Worker error (#{consecutive_errors}): {e}")

                if consecutive_errors > 10:
                    logger.warning("[SENDER] Too many errors, cooling down 30s")
                    await asyncio.sleep(30)
                    consecutive_errors = 0